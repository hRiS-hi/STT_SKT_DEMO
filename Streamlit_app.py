import streamlit as st
import torch
import torchaudio
import io
import numpy as np
import librosa
from transformers import AutoModel
from audio_recorder_streamlit import audio_recorder


#for pdf and docx generation
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt



from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import os

def generate_pdf(title, questions):
    filename = f"{title.replace(' ', '_')}.pdf"

    # ✅ Correct font loading
    FONT_PATH = os.path.join(os.path.dirname(__file__), "NotoSansDevanagari-Regular.ttf")
    pdfmetrics.registerFont(TTFont("Deva", FONT_PATH))

    c = canvas.Canvas(filename, pagesize=A4)
    width, height = A4

    # Header - Use Helvetica-Bold for English

    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width/2, height - 1.4*inch, "ANNUAL EXAMINATION")

    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width/2, height - 1.8*inch, "SUBJECT : SANSKRIT")

    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width/2, height - 2.3*inch, title)

    c.line(1*inch, height - 2.5*inch, width - 1*inch, height - 2.5*inch)

    # Use Devanagari font for Sanskrit text
    c.setFont("Deva", 13)
    c.drawString(1*inch, height - 3*inch, "निम्नलिखित प्रश्नों के उत्तर लिखिए।")

    y = height - 3.7*inch
    c.setFont("Deva", 14)

    for i, q in enumerate(questions):
        text = c.beginText(1*inch, y)
        text.setFont("Deva", 14)
        text.textLines(f"प्रश्न {i+1}: {q}")
        c.drawText(text)

        y -= 1*inch

        if y < 1.2*inch:
            c.showPage()
            c.setFont("Deva", 14)
            y = height - 1.5*inch

    c.save()
    return filename

def generate_docx(title, questions):
    filename = f"{title.replace(' ', '_')}.docx"
    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    title_para.alignment = 1  # center

    doc.add_paragraph("")  # spacing

    # Instruction
    instr = doc.add_paragraph("Answer the following questions:")
    instr.runs[0].font.size = Pt(12)

    doc.add_paragraph("")

    # Questions
    for i, q in enumerate(questions):
        para = doc.add_paragraph(f"Q{i+1}) {q}")
        para.runs[0].font.size = Pt(12)

    doc.save(filename)
    return filename





# --- Configuration & Constants ---
st.set_page_config(page_title="Voice Question Paper Generator", layout="centered")
TARGET_SR = 16000

# --- Session State Initialization ---
if 'question_data' not in st.session_state:
    # Stores the text for each question index: {0: "Question 1 text", 1: "Question 2 text"}
    st.session_state.question_data = {} 

# --- Model Loading (Cached) ---
@st.cache_resource
def load_model():
    """Loads the AI model once and caches it."""
    with st.spinner("Loading Neural Engine... (takes a minute first time)"):
        # Using the multilingual indic conformer
        model = AutoModel.from_pretrained("ai4bharat/indic-conformer-600m-multilingual", trust_remote_code=True)
    return model

# --- Audio Processing ---
def process_audio_data(audio_bytes, target_sr=TARGET_SR):
    """Converts raw audio bytes to standard torch tensor format."""
    try:
        audio_buffer = io.BytesIO(audio_bytes)
        # Try standard torchaudio load first
        wav, sr = torchaudio.load(audio_buffer)
        
        # Convert to mono if needed
        if wav.shape[0] > 1:
             wav = torch.mean(wav, dim=0, keepdim=True)
             
        # Resample to 16kHz if needed
        if sr != target_sr:
            resampler = torchaudio.transforms.Resample(orig_freq=sr, new_freq=target_sr)
            wav = resampler(wav)
            
        return wav
    except Exception:
        # Fallback to librosa for robustness
        audio_buffer.seek(0)
        y, sr = librosa.load(audio_buffer, sr=target_sr, mono=True)
        return torch.tensor(y).unsqueeze(0)

# --- Main App Interface ---
def main():
    st.title("📝 Voice-to-Text Question Paper")
    st.markdown("Speak your questions in Sanskrit (or other Indic languages) to generate a standardized question paper.")
    
    model = load_model()

    # --- Sidebar Settings ---
    with st.sidebar:
        st.header("⚙️ Settings")
        
        # Language Selection (Defaults to Sanskrit 'sa')
        lang_options = {"Sanskrit": "sa", "Hindi": "hi", "Malayalam": "ml", "Tamil": "ta", "Telugu": "te", "Kannada": "kn", "Marathi": "mr", "Gujarati": "gu", "Bengali": "bn", "English": "en"}
        # Set default index to 0 (Sanskrit)
        selected_lang = st.selectbox("Input Language", list(lang_options.keys()), index=0)
        lang_code = lang_options[selected_lang]
        
        st.markdown("---")
        st.markdown("**Instructions:**")
        st.markdown("1. Set total number of questions.")
        st.markdown("2. Click the microphone for a question.")
        st.markdown("3. Speak clearly and stop recording.")
        st.markdown("4. Review and edit the transcribed text if needed.")

    # --- Paper Setup ---
    col_setup_1, col_setup_2 = st.columns([3, 1])
    with col_setup_1:
        paper_title = st.text_input("Paper Title", value="Sanskrit Examination - Term I")
    with col_setup_2:
        num_questions = st.number_input("Total Questions", min_value=1, max_value=50, value=5, step=1)

    st.divider()

    # --- Question Loop ---
    for i in range(num_questions):
        # Visual container for each question
        with st.expander(f"Question {i+1}", expanded=True):
            col_rec, col_text = st.columns([1, 4])
            
            with col_rec:
                st.caption("Record")
                # Unique key for each recorder is CRITICAL
                audio_bytes = audio_recorder(
                    pause_threshold=2.0,
                    sample_rate=TARGET_SR,
                    key=f"rec_{i}",
                    icon_size="2x",
                )

            # If new audio is recorded for this specific question index
            if audio_bytes:
                # Generate a unique hash per audio capture to prevent re-transcribing same clip
                audio_hash = hash(audio_bytes)

                # Initialize placeholders in session state
                if "question_data" not in st.session_state:
                    st.session_state.question_data = {}
                if f"last_audio_{i}" not in st.session_state:
                    st.session_state[f"last_audio_{i}"] = None

                # Only transcribe if new audio was recorded (avoid infinite reruns)
                if st.session_state[f"last_audio_{i}"] != audio_hash:
                    with col_text:
                        with st.spinner("🎙️ Transcribing..."):
                            wav = process_audio_data(audio_bytes)
                            transcription = model(wav, lang_code, "rnnt").strip()

                            # Save transcription and hash
                            st.session_state.question_data[i] = transcription
                            st.session_state[f"last_audio_{i}"] = audio_hash

            # Always read the current transcription (if available)
            # Ensure a separate session key to store widget text
            text_key = f"text_area_{i}"

            # If we have new transcription and widget doesn't have text yet → initialize widget state
            if text_key not in st.session_state:
                st.session_state[text_key] = st.session_state.question_data.get(i, "")

            # If transcription updated session value → push it into widget state
            elif st.session_state.question_data.get(i, "") != st.session_state[text_key]:
                st.session_state[text_key] = st.session_state.question_data[i]

            # Render text area (NO value= parameter!)
            new_text = st.text_area(
                f"Question {i+1} Text",
                key=text_key,
                height=70,
                label_visibility="collapsed",
                placeholder="Speak to transcribe, or type question here..."
            )

            # Save manual edits back to question storage
            st.session_state.question_data[i] = new_text


    # --- Final Output Generation ---
    st.divider()
    st.subheader("📄 Final Paper Preview")
    
    if st.button("Generate Final Paper"):
        questions = [st.session_state.question_data.get(i, "[No question]") for i in range(num_questions)]

        final_text = "\n".join([f"Q{i+1}) {q}" for i, q in enumerate(questions)])

        st.markdown("### 📄 Preview")
        st.text(final_text)

        # ----- PDF -----
        pdf_file = generate_pdf(paper_title, questions)
        with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF (NCERT Style)", f, file_name=pdf_file, mime="application/pdf")

        # ----- WORD -----
        docx_file = generate_docx(paper_title, questions)
        with open(docx_file, "rb") as f:
            st.download_button("⬇️ Download Word (.docx)", f, file_name=docx_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    main()